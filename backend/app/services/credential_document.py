"""Pembuatan register kredensial proyek dalam format Word."""
from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = RGBColor(11, 37, 69)
CYAN = RGBColor(0, 153, 204)
MUTED = RGBColor(91, 107, 128)
TABLE_WIDTHS_DXA = [500, 1650, 2500, 1500, 1350, 1860]


def _set_run(run, *, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _shade(element, fill: str) -> None:
    properties = element.get_or_add_tcPr() if element.tag.endswith("}tc") else element.get_or_add_pPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in TABLE_WIDTHS_DXA:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, TABLE_WIDTHS_DXA[index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def _add_credentials_table(document, rows: list[dict], *, start_number: int) -> None:
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ("No.", "Nama", "Email", "Peran Proyek", "Divisi", "Password Awal")
    for index, label in enumerate(headers):
        cell = table.rows[0].cells[index]
        _shade(cell._tc, "E8EEF5")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        _set_run(paragraph.add_run(label), size=8.5, bold=True, color=NAVY)

    for number, item in enumerate(rows, start=start_number):
        values = (
            str(number),
            item["name"],
            item["email"],
            item["project_role"],
            item.get("division") or "-",
            item["password"],
        )
        cells = table.add_row().cells
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index in (0, 5) else WD_ALIGN_PARAGRAPH.LEFT
            _set_run(paragraph.add_run(value), size=8.2, bold=index == 5, color=NAVY if index == 5 else None)
    _set_table_geometry(table)


def _add_masthead(document) -> None:
    masthead = document.add_paragraph()
    masthead.paragraph_format.space_after = Pt(10)
    _set_run(masthead.add_run("RENCANIX  |  PROJECT CREDENTIAL REGISTER"), size=8.5, bold=True, color=MUTED)


def build_project_credentials_docx(
    *,
    project_name: str,
    admin_name: str,
    rows: list[dict],
    generated_at: datetime | None = None,
) -> bytes:
    generated_at = generated_at or datetime.utcnow()
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    section.header.paragraphs[0].clear()

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("RAHASIA - Simpan terbatas pada Admin Proyek dan pengguna terkait")
    _set_run(footer_run, size=8, color=MUTED)

    _add_masthead(document)
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    _set_run(kicker.add_run("ADMINISTRASI AKUN PROYEK"), size=9, bold=True, color=CYAN)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    _set_run(title.add_run("Daftar Akun dan Password Awal"), size=22, bold=True, color=NAVY)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    _set_run(subtitle.add_run(project_name), size=12, bold=True, color=MUTED)

    for label, value in (
        ("Dibuat oleh", admin_name),
        ("Waktu pembuatan", generated_at.strftime("%d-%m-%Y %H:%M UTC")),
        ("Jumlah akun", str(len(rows))),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        _set_run(paragraph.add_run(f"{label}: "), size=9.5, bold=True, color=NAVY)
        _set_run(paragraph.add_run(value), size=9.5, color=MUTED)

    warning = document.add_paragraph()
    warning.paragraph_format.space_before = Pt(8)
    warning.paragraph_format.space_after = Pt(12)
    _shade(warning._p, "FFF4D6")
    _set_run(
        warning.add_run(
            "Penting: pembuatan dokumen ini merotasi password akun. Password berikut aktif untuk login dan harus diganti oleh setiap pengguna setelah akses pertama."
        ),
        size=9,
        bold=True,
        color=RGBColor(122, 90, 0),
    )

    first_page_count = 14
    next_page_count = 16
    position = 0
    page_capacity = first_page_count
    while position < len(rows):
        if position:
            document.add_page_break()
            _add_masthead(document)
        chunk = rows[position:position + page_capacity]
        _add_credentials_table(document, chunk, start_number=position + 1)
        position += len(chunk)
        page_capacity = next_page_count

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    _set_run(note.add_run("Catatan keamanan: "), size=8.5, bold=True, color=NAVY)
    _set_run(note.add_run("dokumen tidak menyimpan hash dan tidak dapat dibuat ulang dengan password yang sama."), size=8.5, color=MUTED)

    document.core_properties.title = f"Daftar Akun - {project_name}"
    document.core_properties.subject = "Register kredensial awal proyek"
    document.core_properties.author = "Rencanix"
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
