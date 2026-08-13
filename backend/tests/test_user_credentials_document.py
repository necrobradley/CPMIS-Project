import asyncio
import io
import threading
import time

import pytest
from docx import Document
from fastapi import HTTPException, UploadFile
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app.api.v1.endpoints.users import (
    AIEmployeeImportCommitRequest,
    AIEmployeeImportRow,
    CredentialDocumentRequest,
    download_project_credentials_document,
    import_ai_reviewed_employees,
    import_users_from_csv,
    preview_employee_position_mapping,
)
from app.api.v1.endpoints import users as users_endpoint
from app.core.security import get_password_hash, verify_password
from app.db.database import Base
from app.models.user import Division, Project, ProjectMembership, User, UserRole
from app.services.ai_service import AIService


def build_database():
    engine = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    return sessionmaker(bind=engine)()


def seed_project(db):
    admin = User(
        name="Admin Proyek Demo",
        email="admin.project@example.com",
        password_hash=get_password_hash("Admin-Password-123!"),
        role=UserRole.ADMIN,
        is_active=True,
    )
    member = User(
        name="Site Engineer Demo",
        email="site.engineer@example.com",
        password_hash=get_password_hash("Old-Password-123!"),
        role=UserRole.STAFF,
        is_active=True,
        must_set_password=True,
        email_verification_required=True,
    )
    db.add_all([admin, member])
    db.flush()
    project = Project(project_name="Pusat Inovasi Demo", owner_id=admin.id)
    db.add(project)
    db.flush()
    division = Division(project_id=project.id, division_name="Engineering")
    db.add(division)
    db.flush()
    db.add_all([
        ProjectMembership(
            project_id=project.id,
            user_id=admin.id,
            division_id=division.id,
            project_role="project_admin",
        ),
        ProjectMembership(
            project_id=project.id,
            user_id=member.id,
            division_id=division.id,
            project_role="site_engineer",
        ),
    ])
    db.commit()
    return admin, member


async def response_body(response) -> bytes:
    return b"".join([chunk async for chunk in response.body_iterator])


def test_credentials_document_rotates_password_and_contains_login_table():
    db = build_database()
    admin, member = seed_project(db)

    response = download_project_credentials_document(
        CredentialDocumentRequest(
            current_password="Admin-Password-123!",
            confirmation="GENERATE PASSWORD",
        ),
        db=db,
        current_user=admin,
    )
    content = asyncio.run(response_body(response))
    document = Document(io.BytesIO(content))

    assert response.headers["cache-control"] == "no-store, max-age=0"
    assert document.core_properties.title == "Daftar Akun - Pusat Inovasi Demo"
    assert len(document.tables) == 1
    values = [cell.text for cell in document.tables[0].rows[1].cells]
    assert values[1:5] == [
        "Site Engineer Demo",
        "site.engineer@example.com",
        "Site Engineer",
        "Engineering",
    ]
    generated_password = values[5]
    assert len(generated_password) == 16

    db.refresh(member)
    assert verify_password(generated_password, member.password_hash)
    assert not verify_password("Old-Password-123!", member.password_hash)
    assert member.email_verified_at is not None
    assert member.must_set_password is False
    assert member.auth_version == 2


def test_credentials_document_requires_exact_confirmation_without_rotation():
    db = build_database()
    admin, member = seed_project(db)
    old_hash = member.password_hash

    with pytest.raises(HTTPException) as exc:
        download_project_credentials_document(
            CredentialDocumentRequest(
                current_password="Admin-Password-123!",
                confirmation="generate password",
            ),
            db=db,
            current_user=admin,
        )

    assert exc.value.status_code == 400
    db.refresh(member)
    assert member.password_hash == old_hash


def test_temporary_password_hashing_runs_in_parallel(monkeypatch):
    active = 0
    highest_active = 0
    lock = threading.Lock()

    def deliberately_slow_hash(password: str) -> str:
        nonlocal active, highest_active
        with lock:
            active += 1
            highest_active = max(highest_active, active)
        time.sleep(0.04)
        with lock:
            active -= 1
        return f"hashed:{password}"

    monkeypatch.setattr(users_endpoint, "get_password_hash", deliberately_slow_hash)
    started = time.perf_counter()
    credentials = users_endpoint._hashed_temporary_passwords(12)
    elapsed = time.perf_counter() - started

    assert len(credentials) == 12
    assert highest_active > 1
    assert elapsed < 0.30


def test_account_dataset_can_create_project_division_by_name():
    db = build_database()
    admin, _ = seed_project(db)
    csv_content = (
        "name,email,role,phone,telegram_id,division_name,project_role\n"
        "QA Engineer CSV,qa.csv@example.com,staff,,,QA/QC,qa_qc_engineer\n"
    ).encode()
    upload = UploadFile(filename="accounts.csv", file=io.BytesIO(csv_content))

    result = asyncio.run(import_users_from_csv(file=upload, db=db, current_user=admin))

    assert result["created"] == 1
    division = db.query(Division).filter(Division.division_name == "QA/QC").one()
    imported = db.query(User).filter(User.email == "qa.csv@example.com").one()
    membership = db.query(ProjectMembership).filter(ProjectMembership.user_id == imported.id).one()
    assert membership.division_id == division.id
    assert membership.project_role == "qa_qc_engineer"


def test_ai_preview_maps_position_without_creating_account(monkeypatch):
    db = build_database()
    admin, _ = seed_project(db)

    async def fake_mapping(self, employees, available_roles):
        assert employees == [{"row": 2, "position": "Manajer Proyek"}]
        assert all("email" not in employee and "name" not in employee for employee in employees)
        return [{
            "row": 2,
            "project_role": "project_manager",
            "division_name": "Project Management",
            "confidence": 0.97,
            "reasoning": "Posisi memimpin koordinasi proyek.",
        }]

    monkeypatch.setattr(AIService, "is_configured", classmethod(lambda cls, route="default": True))
    monkeypatch.setattr(AIService, "map_employee_positions", fake_mapping)
    monkeypatch.setattr(
        AIService,
        "_route_config",
        classmethod(lambda cls, route="default", provider=None, model=None: {
            "provider": "mlapi", "model": "nemotron-3-ultra",
        }),
    )
    upload = UploadFile(
        filename="employees.csv",
        file=io.BytesIO("name;email;position\nRina;rina@example.com;Manajer Proyek\n".encode()),
    )

    result = asyncio.run(preview_employee_position_mapping(file=upload, db=db, current_user=admin))

    assert result["accounts_created"] == 0
    assert result["ai_provider"] == "mlapi"
    assert result["ai_model"] == "nemotron-3-ultra"
    assert result["ai_status"] == "success"
    assert result["rows"][0]["project_role"] == "project_manager"
    assert result["rows"][0]["role"] == "manager"
    assert db.query(User).filter(User.email == "rina@example.com").first() is None


def test_ai_preview_falls_back_transparently_when_model_is_unavailable(monkeypatch):
    db = build_database()
    admin, _ = seed_project(db)

    async def unavailable_mapping(self, employees, available_roles):
        raise ValueError("model_not_found")

    monkeypatch.setattr(AIService, "is_configured", classmethod(lambda cls, route="default": True))
    monkeypatch.setattr(AIService, "map_employee_positions", unavailable_mapping)
    monkeypatch.setattr(
        AIService,
        "_route_config",
        classmethod(lambda cls, route="default", provider=None, model=None: {
            "provider": "mlapi", "model": "nemotron-3-ultra",
        }),
    )
    upload = UploadFile(
        filename="employees.csv",
        file=io.BytesIO("name,email,position\nRina,rina@example.com,Manajer Proyek\n".encode()),
    )

    result = asyncio.run(preview_employee_position_mapping(file=upload, db=db, current_user=admin))

    assert result["ai_status"] == "fallback"
    assert result["fallback_count"] == 1
    assert result["rows"][0]["project_role"] == "project_manager"
    assert result["rows"][0]["mapping_source"] == "system_fallback"
    assert "model_not_found" not in result["ai_error"]


def test_ai_review_commit_creates_account_with_system_derived_role():
    db = build_database()
    admin, _ = seed_project(db)
    request = AIEmployeeImportCommitRequest(rows=[AIEmployeeImportRow(
        row=2,
        name="Rina",
        email="rina@example.com",
        position="Manajer Proyek",
        role=UserRole.STAFF,
        division_name="Project Management",
        project_role="project_manager",
        confidence=0.97,
        reasoning="Pemetaan telah ditinjau.",
    )])

    result = asyncio.run(import_ai_reviewed_employees(data=request, db=db, current_user=admin))

    assert result["created"] == 1
    assert result["ai_reviewed"] is True
    created = db.query(User).filter(User.email == "rina@example.com").one()
    assert created.role == UserRole.MANAGER
    membership = db.query(ProjectMembership).filter(ProjectMembership.user_id == created.id).one()
    assert membership.project_role == "project_manager"
    assert membership.division.division_name == "Project Management"


def test_ai_position_prompt_excludes_employee_identity(monkeypatch):
    captured = {}

    async def fake_completion(self, system_prompt, user_message, route="default", provider=None, model=None):
        captured["message"] = user_message
        return '[{"row":2,"project_role":"site_engineer","division_name":"Engineering","confidence":0.9,"reasoning":"Teknis lapangan"}]'

    monkeypatch.setattr(AIService, "_chat_completion", fake_completion)
    result = asyncio.run(AIService().map_employee_positions(
        [{"row": 2, "position": "Site Engineer"}],
        [{
            "code": "site_engineer",
            "label": "Site Engineer",
            "default_division": "Engineering",
            "responsibility": "Pekerjaan teknis lapangan",
        }],
    ))

    assert result[0]["project_role"] == "site_engineer"
    assert "rina@example.com" not in captured["message"]
    assert '"name"' not in captured["message"]
    assert '"email"' not in captured["message"]
